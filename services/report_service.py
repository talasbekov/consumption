from sqlalchemy.orm import Session
from sqlalchemy.sql import func
from datetime import date, datetime, timedelta # Added datetime, timedelta
from typing import Optional, List, Dict, Any # Added Any

from fastapi import HTTPException

from models import (
    Employee,
    Division,
    Position, # May not be used if focusing on Employee.division_id
    Status,
    EmployeeStatus,
    SecondmentLog,
    User,
    SecondmentStatusEnum, # For seconded-in calculation
    DivisionTypeEnum # For hierarchy logic
)
# from models import State # Explicitly not importing State

# Placeholder for create_word_report_from_template if it's moved or adapted later
# For now, this service will focus on data generation.
import os # For checking template path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# from utils.report_generator import create_word_report_from_template # This is where it was previously

class ReportService:
    def _format_date(self, date_value: Any) -> str:
        if date_value is None:
            return ""
        elif isinstance(date_value, datetime):
            return date_value.strftime("%d.%m.%Y")
        elif isinstance(date_value, date):
            return date_value.strftime("%d.%m.%Y")
        elif isinstance(date_value, str):
            try:
                parsed_date = datetime.strptime(date_value, "%Y-%m-%d")
                return parsed_date.strftime("%d.%m.%Y")
            except ValueError:
                # Attempt another common format if the first one fails
                try:
                    parsed_date = datetime.strptime(date_value, "%d.%m.%Y")
                    return parsed_date.strftime("%d.%m.%Y") # Or return as is if already in desired format
                except ValueError:
                    # print(f"Ошибка преобразования даты: {date_value}. Неизвестный формат.") # Consider logging
                    return date_value # Return original if all parsing fails
        else:
            # print(f"Неожиданный тип данных для даты: {date_value} ({type(date_value)})") # Consider logging
            return ""

    # Copied from services/state.py and will be adapted
    def get_count_of_state(self, db: Session, user_id: int)-> dict[
        str | Any, dict[str, int | str] | dict[str, int | str] | dict[str, int | str] | dict[
            str, int | str] | dict[str, int | Any]]: # Changed InstrumentedAttribute to Any
        # This method needs significant refactoring to not use self.model and user_id for department context
        # For now, porting as is and will refactor in next step.
        print(user_id)
        # user = user_service.get_by_id(db, user_id) # user_service is not defined here yet
        # print(user.employee_id)
        emps = []
        # states = db.query(self.model).all() # self.model is State, needs to change
        # for state_obj in states: # Renamed state to state_obj to avoid conflict with local var
        #     emps.append(state_obj.employee_id)

        # state = None # This variable 'state' refers to an instance of the old State model.
        index = 0

        # while state is None and index < len(emps):
        #     employee_id = emps[index]
        #     state = db.query(self.model).filter(self.model.employee_id == employee_id).first()
        #     index += 1

        # if state is not None:
        #     print("Found state:", state)
        # else:
        #     print("No valid state found")
        #     # This will cause errors below if state is None, needs guard or proper context
        #     # For now, returning empty or raising error if context cannot be determined
        #     raise HTTPException(status_code=404, detail="Reporting context (e.g., department) could not be determined for the user.")


        # state_count = db.query(self.model).filter(self.model.department_id == state.department_id).count()
        # vacant_count = db.query(self.model).filter(self.model.employee_id.is_(None)).count()
        # by_list_count = state_count - vacant_count
        # inline_count = (
        #     db.query(Employee)
        #     .join(EmployeeStatus, EmployeeStatus.employee_id == Employee.id)
        #     .join(Status, Status.id == EmployeeStatus.status_id)
        #     .join(self.model, self.model.employee_id == Employee.id) # self.model is State
        #     .filter(self.model.department_id == state.department_id)
        #     .filter(Status.nameRU == "в строю") # TODO: Change to Status.code
        #     .count()
        # )
        # by_status_count = by_list_count - inline_count

        # result = {
        #     "by state": {"count": state_count, "name": "по штату"},
        #     "by list": {"count": by_list_count, "name": "по списку"},
        #     "by vacant": {"count": vacant_count, "name": "ваканты"},
        #     "by status": {"count": by_status_count, "name": "общее количество отсутствующих"},
        # }

        # statuses = db.query(Status).all()
        # for status_obj in statuses: # Renamed status to status_obj
        #     count = (
        #         db.query(Employee)
        #         .join(EmployeeStatus, EmployeeStatus.employee_id == Employee.id)
        #         .join(Status, Status.id == EmployeeStatus.status_id)
        #         .join(self.model, self.model.employee_id == Employee.id) # self.model is State
        #         .filter(self.model.department_id == state.department_id)
        #         .filter(Status.nameRU == status_obj.nameRU) # TODO: Change to Status.code
        #         .count()
        #     )
        #     result[status_obj.nameEN] = {"count": count, "name": status_obj.nameRU}
        # return result
        raise NotImplementedError("get_count_of_state has been superseded by generate_division_report_data.")


    # Copied from services/state.py and will be adapted
    def generate_division_report_data(self, db: Session, division_id: int, report_date: date) -> Dict[str, Any]:
        main_division = db.query(Division).filter(Division.id == division_id).first()
        if not main_division:
            raise HTTPException(status_code=404, detail=f"Division with id {division_id} not found.")

        # 1. Hierarchy Traversal
        sub_divisions_to_report: List[Division] = []
        if main_division.division_type == DivisionTypeEnum.COMPANY:
            # Example: Report on Departments within a Company
            children = db.query(Division).filter(
                Division.parent_division_id == main_division.id,
                Division.division_type == DivisionTypeEnum.DEPARTMENT
            ).all()
            sub_divisions_to_report.extend(children)
        elif main_division.division_type == DivisionTypeEnum.DEPARTMENT:
            # Example: Report on Directorates (Управления) and direct Offices (Отделы) within a Department
            # Add Department itself if it can have employees directly for the report structure
            # sub_divisions_to_report.append(main_division) # Uncomment if Department itself is a report line
            children_directorates = db.query(Division).filter(
                Division.parent_division_id == main_division.id,
                Division.division_type == DivisionTypeEnum.DIRECTORATE
            ).all()
            sub_divisions_to_report.extend(children_directorates)
            children_offices = db.query(Division).filter(
                Division.parent_division_id == main_division.id,
                Division.division_type == DivisionTypeEnum.DIVISION # Assuming DIVISION is Отдел
            ).all()
            sub_divisions_to_report.extend(children_offices)
        elif main_division.division_type == DivisionTypeEnum.DIRECTORATE: # Управление
             # Add Directorate itself if it can have employees directly
            # sub_divisions_to_report.append(main_division) # Uncomment if Directorate itself is a report line
            children_offices = db.query(Division).filter(
                Division.parent_division_id == main_division.id,
                Division.division_type == DivisionTypeEnum.DIVISION
            ).all()
            sub_divisions_to_report.extend(children_offices)
        elif main_division.division_type == DivisionTypeEnum.DIVISION: # Отдел
            sub_divisions_to_report.append(main_division) # Report on the Отдел itself

        if not sub_divisions_to_report and main_division: # If no children found by type, report on the main_division itself
            sub_divisions_to_report.append(main_division)

        # Ensure unique divisions if a parent could also be listed via children logic (depends on exact needs)
        unique_sub_divisions_dict = {div.id: div for div in sub_divisions_to_report}
        sub_divisions_to_report = list(unique_sub_divisions_dict.values())


        statuses_db = db.query(Status).all()
        status_info_by_code: Dict[str, Dict[str, Any]] = {
            s.code: {"id": s.id, "nameRU": s.nameRU, "nameEN": s.nameEN, "is_active_list_status": True} # Default to true
            for s in statuses_db if s.code
        }
        # Example: Define which statuses are NOT part of "Список" if applicable
        # if 'SECONDED_OUT' in status_info_by_code: status_info_by_code['SECONDED_OUT']['is_active_list_status'] = False


        report_table_rows: List[Dict[str, Any]] = []

        for current_report_division in sub_divisions_to_report:
            division_row_data: Dict[str, Any] = {
                "name": current_report_division.nameRU or f"Division_{current_report_division.id}",
                "status_details": {
                    s_code: {"count": 0, "employees": []} for s_code in status_info_by_code.keys()
                },
                "seconded_in_employees_details": [] # For column 13 FIOs
            }

            employees_in_division = db.query(Employee).filter(Employee.division_id == current_report_division.id).all()

            current_spisok_total_for_division = 0
            for emp in employees_in_division:
                current_employee_status_record = db.query(EmployeeStatus)\
                    .join(Status, Status.id == EmployeeStatus.status_id)\
                    .filter(EmployeeStatus.employee_id == emp.id)\
                    .filter(EmployeeStatus.start_date <= report_date)\
                    .filter((EmployeeStatus.end_date == None) | (EmployeeStatus.end_date >= report_date))\
                    .order_by(EmployeeStatus.start_date.desc(), EmployeeStatus.id.desc())\
                    .first()

                if current_employee_status_record:
                    status_model = db.query(Status).filter(Status.id == current_employee_status_record.status_id).first()
                    if status_model and status_model.code and status_model.code in status_info_by_code:
                        # 2. "Список" (List) Calculation
                        # Count if considered an active list status (e.g., not "FIRED" or similar)
                        # The old report included "Откомандирован" (SECONDED_OUT) in column 12, implying they are part of "Список".
                        # Modify status_info_by_code[status_model.code]['is_active_list_status'] if some shouldn't count.
                        if status_info_by_code[status_model.code]['is_active_list_status']:
                             current_spisok_total_for_division += 1

                        division_row_data["status_details"][status_model.code]["count"] += 1
                        division_row_data["status_details"][status_model.code]["employees"].append({
                            "fio": f"{emp.surname} {emp.firstname[0].upper() if emp.firstname else ''}.",
                            "period": f"{self._format_date(current_employee_status_record.start_date)}-{self._format_date(current_employee_status_record.end_date)}",
                            "comment": current_employee_status_record.note
                        })

            division_row_data["spisok_total"] = current_spisok_total_for_division
            # 4. Placeholders for "Штат" and "Вакантные"
            division_row_data["shtat_total"] = current_spisok_total_for_division  # Placeholder
            division_row_data["vacant_total"] = 0  # Placeholder

            division_row_data["inline_count"] = division_row_data["status_details"].get('IN_SERVICE', {}).get('count', 0)

            # 3. "Прикомандированные" (+N) Calculation
            seconded_in_q = db.query(SecondmentLog)\
                .join(Employee, Employee.id == SecondmentLog.employee_id)\
                .filter(SecondmentLog.target_division_id == current_report_division.id)\
                .filter(SecondmentLog.status == SecondmentStatusEnum.ACTIVE)\
                .filter(SecondmentLog.secondment_start_date <= report_date)\
                .filter((SecondmentLog.actual_end_date == None) | (SecondmentLog.actual_end_date > report_date))

            division_row_data["seconded_in_count"] = seconded_in_q.count()
            for sec_log in seconded_in_q.all():
                emp = sec_log.employee
                division_row_data["seconded_in_employees_details"].append({
                    "fio": f"{emp.surname} {emp.firstname[0].upper() if emp.firstname else ''}.",
                    "period": f"{self._format_date(sec_log.secondment_start_date)}-{self._format_date(sec_log.expected_end_date)}",
                    "comment": f"Из ID {sec_log.original_division_id}"
                })

            report_table_rows.append(division_row_data)

        # 6. "Барлығы" (Totals) Row
        overall_totals: Dict[str, Any] = {
            "name": "Барлығы", "shtat_total": 0, "spisok_total": 0, "vacant_total": 0,
            "inline_count": 0, "seconded_in_count": 0,
            "status_details": {
                s_code: {"count": 0, "employees": []} for s_code in status_info_by_code.keys()
            },
            "seconded_in_employees_details": [] # FIOs not typically summed
        }
        for row_data in report_table_rows:
            overall_totals["shtat_total"] += row_data["shtat_total"]
            overall_totals["spisok_total"] += row_data["spisok_total"]
            overall_totals["vacant_total"] += row_data["vacant_total"]
            overall_totals["inline_count"] += row_data["inline_count"]
            overall_totals["seconded_in_count"] += row_data["seconded_in_count"]
            for s_code, details in row_data["status_details"].items():
                if s_code in overall_totals["status_details"]: # Should always be true due to initialization
                    overall_totals["status_details"][s_code]["count"] += details["count"]

        # 5. Data Structure for Report
        final_report_data = {
            "division_name": main_division.nameRU or f"Division_{main_division.id}",
            "report_date_str": self._format_date(report_date), # Used in template
            "report_table_rows": report_table_rows,
            "totals": overall_totals,
            "comment_shtat_vacant": "Поля 'Штат' и 'Вакантные' являются плейсхолдерами. Их точный расчет требует дополнительной бизнес-логики или данных о штатных единицах."
        }

        return final_report_data


    # def _calculate_overall_totals(self, db: Session, main_division: Division, report_date: date, statuses_db: List[Status]) -> Dict[str, Any]:
    #     # This method's logic has been integrated into generate_division_report_data.
    #     # It can be removed or kept if future complex total calculations are needed.
    #     # For now, commenting out as it's not directly used.
    #     pass

    def create_word_report_from_template(
        self,
        db: Session,
        division_id: int,
        report_date: date, # This is date_from for period logic
        for_period: bool = False,
        date_to: Optional[date] = None
    ) -> BytesIO:

        doc = Document() # Create a new document for each call if not appending for period, or handle outside.
                        # For multi-day reports, a single doc is built up.

        # If for_period is False, date_to is ignored, effectively.
        # If for_period is True but date_to is None, it's like for_period=False for report_date.
        if not for_period or not date_to or date_to < report_date:
            loop_dates = [report_date]
        else:
            loop_dates = []
            current_loop_date = report_date
            while current_loop_date <= date_to:
                loop_dates.append(current_loop_date)
                current_loop_date += timedelta(days=1)

        first_page = True
        for single_report_date in loop_dates:
            if not first_page:
                doc.add_page_break()
            first_page = False

            report_data = self.generate_division_report_data(db, division_id, single_report_date)

            # --- Start filling the document for this single_report_date ---
            # Paragraph for the report title and date
            # Using report_data["report_date_str"] which is already formatted by _format_date
            title_str = f"ЖЕТІНШІ ДЕПАРТАМЕНТ ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ {report_data['report_date_str']} ЖЫЛҒЫ"
            # The original code had "ЖЫЛҒЫ" which means "AS OF YEAR". If it means "AS OF DATE", it should be different.
            # Assuming the original intent was "as of {date}".
            # title_str = f"САПТЫҚ ТІЗІМІ ({report_data['division_name']}) - {report_data['report_date_str']}"

            p_title = doc.add_paragraph()
            run_title = p_title.add_run(title_str)
            run_title.font.name = 'Times New Roman'
            run_title.font.size = Pt(14)
            run_title.bold = True
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add some spacing after title - can be adjusted
            doc.add_paragraph()

            # Table creation - Assuming a structure similar to the old template.
            # Number of columns based on typical report:
            # No | Подразделение | Штат | Список | Вакантные | В строю | ...other statuses... | Прикомандированные (+)
            # Statuses from status_info_by_code will define the dynamic columns

            # Fixed columns: No, Подразделение, Штат, Список, Вакантные, В строю
            # Dynamic status columns (excluding "IN_SERVICE" as it's "В строю")
            # Last column: Прикомандированные (+)

            # Get ordered status codes for columns (ensure consistent order)
            # Example order: IN_SERVICE, ON_DUTY, AFTER_DUTY, BUSINESS_TRIP, ON_STUDYING, ON_LEAVE, ON_SICK_LEAVE, SECONDED_OUT
            # This order needs to match the template's expected column order for statuses.
            # For now, using arbitrary order from status_info_by_code keys, excluding IN_SERVICE.
            # This part is crucial and needs to map to specific columns in template.

            status_codes_for_columns = [
                'IN_SERVICE', 'ON_DUTY', 'ON_MISSION', 'ON_VACATION',
                'ON_SICK_LEAVE', 'SECONDED_OUT'
                # TODO: This list MUST match the template.json or expected columns.
                # Add other relevant status codes in their correct report column order.
                # E.g., ON_STUDYING, AFTER_DUTY etc.
            ]

            # Filter status_info_by_code to only those we will display as separate columns
            display_statuses = {code: report_data["totals"]["status_details"][code] for code in status_codes_for_columns if code in report_data["totals"]["status_details"]}

            num_fixed_cols = 6 # No, Name, Shtat, Spisok, Vacant, V Stroyu
            num_dynamic_status_cols = len(display_statuses) -1 # -1 because IN_SERVICE is already "V Stroyu"
            num_seconded_in_col = 1 # For "Прикомандированные (+)"
            total_cols = num_fixed_cols + num_dynamic_status_cols + num_seconded_in_col

            table = doc.add_table(rows=1, cols=total_cols)
            table.style = 'Table Grid' # Basic grid style

            # Header Row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'Атауы (Наименование)' # Name
            hdr_cells[2].text = 'Штат' # by state
            hdr_cells[3].text = 'Тізім бойынша (По списку)' # by list
            hdr_cells[4].text = 'Бос орындар (Вакант)' # vacant
            hdr_cells[5].text = 'Саптағылар (В строю)' # IN_SERVICE

            col_idx = 6
            for code in status_codes_for_columns:
                if code == 'IN_SERVICE': continue # Already handled by "В строю"
                status_name_ru = db.query(Status.nameRU).filter(Status.code == code).scalar_one_or_none() or code
                hdr_cells[col_idx].text = status_name_ru
                col_idx += 1
            hdr_cells[col_idx].text = 'Іссапарға келгендер (+)' # Seconded In

            # Data Rows
            for idx, row_data in enumerate(report_data["report_table_rows"]):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx + 1)
                row_cells[1].text = row_data["name"]
                row_cells[2].text = str(row_data["shtat_total"])
                row_cells[3].text = str(row_data["spisok_total"])
                row_cells[4].text = str(row_data["vacant_total"])
                row_cells[5].text = str(row_data["inline_count"]) # IN_SERVICE count

                col_idx = 6
                for code in status_codes_for_columns:
                    if code == 'IN_SERVICE': continue
                    row_cells[col_idx].text = str(row_data["status_details"].get(code, {}).get("count", 0))
                    # Add employee details to cell if needed (very small font)
                    # This makes cells very busy. Usually, details are in appendix or separate report.
                    # For now, just counts as per typical high-level report.
                    # Example for details:
                    # emp_details_str = ""
                    # for emp_info in row_data["status_details"].get(code,{}).get("employees",[]):
                    #    emp_details_str += f"{emp_info['fio']} ({emp_info['period']}); "
                    # if emp_details_str:
                    #    run = row_cells[col_idx].paragraphs[0].add_run(f"\n{emp_details_str}")
                    #    run.font.size = Pt(6)
                    col_idx += 1
                row_cells[col_idx].text = str(row_data["seconded_in_count"])
                # Add seconded-in employee details if needed
                # emp_details_sec_in_str = ""
                # for emp_info in row_data.get("seconded_in_employees_details",[]):
                #    emp_details_sec_in_str += f"{emp_info['fio']} ({emp_info['period']}); "
                # if emp_details_sec_in_str:
                #    run = row_cells[col_idx].paragraphs[0].add_run(f"\n{emp_details_sec_in_str}")
                #    run.font.size = Pt(6)


            # Totals Row ("Барлығы")
            totals_data = report_data["totals"]
            total_cells = table.add_row().cells
            total_cells[1].text = totals_data["name"] # "Барлығы"
            total_cells[1].paragraphs[0].runs[0].bold = True # Make "Барлығы" bold
            total_cells[2].text = str(totals_data["shtat_total"])
            total_cells[3].text = str(totals_data["spisok_total"])
            total_cells[4].text = str(totals_data["vacant_total"])
            total_cells[5].text = str(totals_data["inline_count"]) # Total IN_SERVICE

            col_idx = 6
            for code in status_codes_for_columns:
                if code == 'IN_SERVICE': continue
                total_cells[col_idx].text = str(totals_data["status_details"].get(code, {}).get("count", 0))
                col_idx += 1
            total_cells[col_idx].text = str(totals_data["seconded_in_count"])

            # Apply bold to all cells in totals row
            for cell in total_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Comment about placeholders if needed
            doc.add_paragraph(report_data["comment_shtat_vacant"], style='Caption')
            # --- End filling document for this single_report_date ---

        # Save the document to a BytesIO buffer
        file_buffer = BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)

        return file_buffer

# This instantiation will be done in services/__init__.py as per instructions
