import logging
import os
from datetime import datetime, date
from typing import Optional, List
from fastapi import HTTPException

from sqlalchemy import func
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.orm import Session, InstrumentedAttribute

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

from models import State, Employee, Status, Management, EmployeeStatus
from schemas import (
    StateCreate,
    StateUpdate,
)  # Предполагается, что у вас есть схемы создания и обновления событий
from schemas import EmployeeDataBulkUpdate
from services import user_service

from services.base import ServiceBase

'''[
  {
    "employee_id": 8,
    "rank_id": 5,
    "sort": 88,
    "note": "Test",
    "statuses": {
      "employee_id": 8,
      "status_id": 4,
      "start_date": "2025-01-21",
      "end_date": "2025-02-21",
      "note": "Test"
    }
  }
]'''


class StateService(ServiceBase[State, StateCreate, StateUpdate]):

    def get_by_employee_id(self, db: Session, employee_id: int) -> Optional[State]:
        return db.query(State).filter(State.employee_id == employee_id).first()

    async def update_employees_by_state(self, db: Session, employees_data: List[EmployeeDataBulkUpdate]):
        """
        Обновляет данные сотрудников на основе входящего списка данных, включая статусы.
        """
        try:
            employee_ids = [employee_data.employee_id for employee_data in employees_data if employee_data.employee_id]

            # Получаем всех сотрудников одним запросом
            employees = db.query(Employee).filter(Employee.id.in_(employee_ids)).all()
            employees_dict = {employee.id: employee for employee in employees}

            for employee_data in employees_data:
                employee = employees_dict.get(employee_data.employee_id)
                if not employee:
                    logging.warning(f"Employee with ID {employee_data.employee_id} not found")
                    continue

                # Обновляем основные данные сотрудника
                if employee_data.sort is not None:
                    employee.sort = employee_data.sort
                if employee_data.rank_id is not None:
                    employee.rank_id = employee_data.rank_id
                if employee_data.note is not None:
                    employee.note = employee_data.note

                # for chatgpt
                employee_status = db.query(EmployeeStatus).filter(
                    EmployeeStatus.employee_id == employee_data.employee_id).first()

                if not employee_status:
                    # Create a new EmployeeStatus if it does not exist
                    emp_status = EmployeeStatus(
                        employee_id=employee_data.employee_id,
                        status_id=employee_data.statuses.status_id,
                        start_date=employee_data.statuses.start_date,
                        end_date=employee_data.statuses.end_date,
                        note=employee_data.statuses.note if employee_data.statuses.note else "",
                    )
                    db.add(emp_status)
                else:
                    # Update the existing EmployeeStatus
                    status_data = employee_data.statuses

                    if status_data:
                        # Update only if the new status data is provided
                        employee_status.status_id = status_data.status_id
                        employee_status.start_date = status_data.start_date
                        employee_status.end_date = status_data.end_date
                        employee_status.note = status_data.note

                    logging.info(f"Status updated for Employee ID {employee_data.employee_id}")

            # Сохраняем изменения
            db.commit()
            logging.info("All employees updated successfully")
        except SQLAlchemyError as e:
            logging.error(f"Database error while processing employees: {e}")
            db.rollback()
            raise HTTPException(status_code=500, detail="An error occurred while updating employees")
        except Exception as e:
            logging.error(f"Unexpected error: {e}")
            db.rollback()
            raise HTTPException(status_code=500, detail="An unexpected error occurred")

    def get_count_of_state(self, db: Session, user_id: int)-> dict[
        str | InstrumentedAttribute, dict[str, int | str] | dict[str, int | str] | dict[str, int | str] | dict[
            str, int | str] | dict[str, int | InstrumentedAttribute]]:
        print(user_id)
        user = user_service.get_by_id(db, user_id)
        print(user.employee_id)
        emps = []
        states = db.query(self.model).all()
        for state in states:
            emps.append(state.employee_id)

        # Инициализируем state как None и создаем итератор по emps
        state = None
        index = 0

        # Используем цикл while для перебора emps
        while state is None and index < len(emps):
            employee_id = emps[index]
            state = db.query(self.model).filter(self.model.employee_id == employee_id).first()
            index += 1

        # Теперь state содержит либо значение, либо остался None, если подходящего значения не нашлось
        if state is not None:
            print("Found state:", state)
        else:
            print("No valid state found")

        state_count = db.query(self.model).filter(self.model.department_id == state.department_id).count()
        vacant_count = db.query(self.model).filter(self.model.employee_id.is_(None)).count()
        by_list_count = state_count - vacant_count
        inline_count = (
            db.query(Employee)
            .join(EmployeeStatus, EmployeeStatus.employee_id == Employee.id)  # Присоединяем таблицу Status через отношение
            .join(Status, Status.id == EmployeeStatus.status_id)
            .join(State, State.employee_id == Employee.id)  # Присоединяем State по employee_id
            .filter(State.department_id == state.department_id)  # Фильтруем по department_id
            .filter(Status.nameRU == "в строю")  # Условие для статуса
            .count()
        )
        by_status_count = by_list_count - inline_count

        # Формируем начальный словарь
        result = {
            "by state": {"count": state_count, "name": "по штату"},
            "by list": {"count": by_list_count, "name": "по списку"},
            "by vacant": {"count": vacant_count, "name": "ваканты"},
            "by status": {"count": by_status_count, "name": "общее количество отсутствующих"},
        }

        # Обрабатываем статусы
        statuses = db.query(Status).all()
        for status in statuses:
            count = (
                db.query(Employee)
                .join(EmployeeStatus, EmployeeStatus.employee_id == Employee.id)  # Присоединяем таблицу Status через отношение
                .join(Status, Status.id == EmployeeStatus.status_id)
                .join(State, State.employee_id == Employee.id)  # Присоединяем State по employee_id
                .filter(State.department_id == state.department_id)  # Фильтруем по department_id
                .filter(Status.nameRU == status.nameRU)  # Фильтруем по имени статуса
                .count()
            )

            # Добавляем данные по каждому статусу в словарь
            result[status.nameEN] = {"count": count, "name": status.nameRU}

        # Возвращаем итоговый словарь
        return result

    def get_count_of_management_state(self, db: Session, user_id: int) -> dict:
        # Получаем пользователя и его состояние
        user = user_service.get_by_id(db, user_id)
        if not user or not user.employee_id:
            raise ValueError("Пользователь не найден или не связан с сотрудником.")

        state = self.get_by_employee_id(db, user.employee_id)
        if not state:
            raise ValueError("Не найдено состояние для пользователя.")

        # Инициализируем итоговый словарь
        result = {}

        # Получаем список всех статусов
        statuses = db.query(Status).all()
        statuses_dict = {status.id: {"nameRU": status.nameRU, "nameEN": status.nameEN} for status in statuses}

        state_count = db.query(func.count(self.model.id)).filter(
            self.model.department_id == state.department_id,
            self.model.division_id.is_(None),
            self.model.management_id.is_(None)
        ).scalar()

        vacant_count = db.query(func.count(self.model.id)).filter(
            self.model.department_id == state.department_id,
            self.model.employee_id.is_(None),
            self.model.division_id.is_(None),
            self.model.management_id.is_(None)
        ).scalar()

        by_list_count = state_count - vacant_count

        # Формируем базовую структуру для управления
        department_data = {
            "by state": {"count": state_count, "name": "по штату"},
            "by list": {"count": by_list_count, "name": "по списку"},
            "by vacant": {"count": vacant_count, "name": "ваканты"},
        }

        # Получаем сотрудников для каждого статуса
        for status_id, status_info in statuses_dict.items():
            employees_with_status = (
                db.query(
                    Employee.surname,
                    Employee.firstname,
                    Status.nameRU,
                    EmployeeStatus.start_date,
                    EmployeeStatus.end_date,
                    EmployeeStatus.note
                )
                .join(EmployeeStatus, EmployeeStatus.employee_id == Employee.id)
                .join(Status, Status.id == EmployeeStatus.status_id)
                .join(State, State.employee_id == Employee.id)
                .filter(
                    State.department_id == state.department_id,
                    State.division_id.is_(None),
                    State.management_id.is_(None),
                    Status.id == status_id
                )
                .all()
            )
            print(employees_with_status, "qqq")

            # Формируем список сотрудников и их данные для текущего статуса
            employees_data = []
            for surname, firstname, nameRU, start_date, end_date, note in employees_with_status:
                def format_date(date_value):
                    if date_value is None:
                        return ""  # Возвращаем пустую строку, если дата отсутствует
                    elif isinstance(date_value, datetime):
                        return date_value.strftime("%d.%m.%Y")  # Форматируем объект datetime
                    elif isinstance(date_value, date):
                        # Преобразуем date в строку формата дд.мм.гггг
                        return date_value.strftime("%d.%m.%Y")
                    elif isinstance(date_value, str):
                        try:
                            # Предполагаем, что строка в формате YYYY-MM-DD
                            parsed_date = datetime.strptime(date_value, "%Y-%m-%d")
                            return parsed_date.strftime("%d.%m.%Y")
                        except ValueError as e:
                            print(f"Ошибка преобразования даты: {date_value}. {e}")  # Логируем ошибку
                            return date_value  # Возвращаем исходное значение, если формат не подходит
                    else:
                        print(
                            f"Неожиданный тип данных для даты: {date_value} ({type(date_value)})")  # Логируем неожиданный тип
                        return ""

                # Пробуем отформатировать даты
                start_date_formatted = format_date(start_date)
                end_date_formatted = format_date(end_date)

                # Логирование для проверки результатов
                print(f"Обработка сотрудника: {surname} {firstname}, {nameRU}")
                print(f"Начало: {start_date} -> {start_date_formatted}")
                print(f"Окончание: {end_date} -> {end_date_formatted}")

                employees_data.append({
                    "ФИО": f"{surname} {firstname[0].upper()}.",
                    "Статус": nameRU,
                    "Дата начало": start_date_formatted,
                    "Дата окончание": end_date_formatted,
                    "Примечание": note
                })

            # Добавляем информацию о статусе в структуру управления
            department_data[status_info["nameEN"]] = {
                "count": len(employees_data),
                "name": f"{status_info['nameRU']} ({status_info['nameEN']})",
                f"сотрудники со статусом {status_info['nameRU']}": employees_data
            }

        # Добавляем данные управления в итоговый результат
        result["Басшылық"] = department_data



        # Получаем все управления текущего департамента
        managements = db.query(Management).filter(Management.department_id == state.department_id).all()
        # Основной цикл по управлениям
        for management in managements:
            # Считаем количество сотрудников "по штату", "по списку" и вакантных
            state_count = db.query(func.count(self.model.id)).filter(
                self.model.department_id == management.department_id,
                self.model.management_id == management.id
            ).scalar()

            vacant_count = db.query(func.count(self.model.id)).filter(
                self.model.employee_id.is_(None),
                self.model.department_id == management.department_id,
                self.model.management_id == management.id
            ).scalar()

            by_list_count = state_count - vacant_count

            # Формируем базовую структуру для управления
            management_data = {
                "by state": {"count": state_count, "name": "по штату"},
                "by list": {"count": by_list_count, "name": "по списку"},
                "by vacant": {"count": vacant_count, "name": "ваканты"},
            }

            # Получаем сотрудников для каждого статуса
            for status_id, status_info in statuses_dict.items():
                employees_with_status = (
                    db.query(
                        Employee.surname,
                        Employee.firstname,
                        Status.nameRU,
                        EmployeeStatus.start_date,
                        EmployeeStatus.end_date,
                        EmployeeStatus.note
                    )
                    .join(EmployeeStatus, EmployeeStatus.employee_id == Employee.id)
                    .join(Status, Status.id == EmployeeStatus.status_id)
                    .join(State, State.employee_id == Employee.id)
                    .filter(
                        State.department_id == management.department_id,
                        State.management_id == management.id,
                        Status.id == status_id
                    )
                    .all()
                )

                # Формируем список сотрудников и их данные для текущего статуса
                employees_data = []
                for surname, firstname, nameRU, start_date, end_date, note in employees_with_status:
                    def format_date(date_value):
                        if date_value is None:
                            return ""  # Возвращаем пустую строку, если дата отсутствует
                        elif isinstance(date_value, datetime):
                            return date_value.strftime("%d.%m.%Y")  # Форматируем объект datetime
                        elif isinstance(date_value, date):
                            # Преобразуем date в строку формата дд.мм.гггг
                            return date_value.strftime("%d.%m.%Y")
                        elif isinstance(date_value, str):
                            try:
                                # Предполагаем, что строка в формате YYYY-MM-DD
                                parsed_date = datetime.strptime(date_value, "%Y-%m-%d")
                                return parsed_date.strftime("%d.%m.%Y")
                            except ValueError as e:
                                print(f"Ошибка преобразования даты: {date_value}. {e}")  # Логируем ошибку
                                return date_value  # Возвращаем исходное значение, если формат не подходит
                        else:
                            print(
                                f"Неожиданный тип данных для даты: {date_value} ({type(date_value)})")  # Логируем неожиданный тип
                            return ""

                    # Пробуем отформатировать даты
                    start_date_formatted = format_date(start_date)
                    end_date_formatted = format_date(end_date)

                    # Логирование для проверки результатов
                    print(f"Обработка сотрудника: {surname} {firstname}, {nameRU}")
                    print(f"Начало: {start_date} -> {start_date_formatted}")
                    print(f"Окончание: {end_date} -> {end_date_formatted}")

                    employees_data.append({
                        "ФИО": f"{surname} {firstname[0].upper()}.",
                        "Статус": nameRU,
                        "Дата начало": start_date_formatted,
                        "Дата окончание": end_date_formatted,
                        "Примечание": note
                    })

                # Добавляем информацию о статусе в структуру управления
                management_data[status_info["nameEN"]] = {
                    "count": len(employees_data),
                    "name": f"{status_info['nameRU']} ({status_info['nameEN']})",
                    f"сотрудники со статусом {status_info['nameRU']}": employees_data
                }

            # Добавляем данные управления в итоговый результат
            result[management.nameKZ] = management_data

        # Получаем общий результат
        overall_data = self.get_count_of_state(db, user_id)

        # Удаляем ключ "by status", если он есть
        if "by status" in overall_data:
            del overall_data["by status"]

        # Сохраняем в result
        result["Барлығы"] = overall_data

        return result

    def create_word_report_from_template(self, db: Session, user_id: int):
        """
        Создание отчета Word с подстановкой данных из функции get_count_of_state и get_count_of_management_state.
        """
        # Получаем данные из функции get_count_of_management_state
        management_data = self.get_count_of_management_state(db, user_id)

        # Получаем текущую дату в формате день.месяц.год
        current_date = datetime.now().strftime("%d.%m.%Y")

        file_path = './template.docx'
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл {file_path} не найден. Пожалуйста, проверьте путь к файлу.")

        doc = Document(file_path)

        # Изменяем первый параграф (перед таблицей)
        first_paragraph = doc.paragraphs[0]  # Получаем первый параграф документа
        run = first_paragraph.add_run(f"ЖЕТІНШІ ДЕПАРТАМЕНТ ЖЕКЕ ҚҰРАМЫНЫҢ САПТЫҚ ТІЗІМІ {current_date} ЖЫЛҒЫ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True  # Делаем текст жирным
        first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Поиск таблицы в документе
        table = doc.tables[0]  # Предположим, что нужная таблица первая

        # Заполнение таблицы данными по каждому управлению
        row_idx = 2  # Строка для первого управления
        total_rows = len(table.rows)

        for management_name, management_info in management_data.items():
            # Берём строку для конкретного управления
            row = table.rows[row_idx]
            # В ячейку с индексом 1 записываем название управления
            row.cells[1].text = management_name

            # Заполняем данные по каждому «основному» статусу для этого управления
            management_values = [
                str(management_info.get("by state", {}).get("count", 0)),
                str(management_info.get("by list", {}).get("count", 0)),
                str(management_info.get("by vacant", {}).get("count", 0)),
                str(management_info.get("inline", {}).get("count", 0)),
                str(management_info.get("on duty", {}).get("count", 0)),
                str(management_info.get("after duty", {}).get("count", 0)),
                str(management_info.get("business trip", {}).get("count", 0)),
                str(management_info.get("on studying", {}).get("count", 0)),
                str(management_info.get("on leave", {}).get("count", 0)),
                str(management_info.get("on sick leave", {}).get("count", 0)),
                str(management_info.get("out seconded", {}).get("count", 0)),
                str(management_info.get("on seconded", {}).get("count", 0))
            ]

            # Записываем основные статусы (упрощённая логика)
            for i, value in enumerate(management_values, start=2):  # Начинаем с 2
                if i >= len(row.cells):
                    # Предохранитель, если в таблице меньше столбцов, чем мы ожидаем
                    break

                cell = row.cells[i]
                cell.text = value

                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_main = paragraph.runs[0]
                run_main.font.size = Pt(12)
                run_main.font.name = 'Times New Roman'

            # Далее ВНУТРЕННИЙ ЦИКЛ по элементам словаря management_info.
            # Раньше здесь использовалась та же переменная i; теперь используем j.
            for j, (status_key, status_info) in enumerate(management_info.items(), start=2):
                # Пропускаем статус "inline", если такова логика
                if status_key == "inline":
                    continue

                # Определяем ячейку, куда писать
                if j >= len(row.cells):
                    # Не пишем, если столбцы закончились
                    break

                cell = row.cells[j]
                # Очищаем ячейку для форматирования
                cell.text = ""

                # Формируем текст для статуса
                paragraph = cell.paragraphs[0]

                # Записываем счётчик (кол-во сотрудников в этом статусе)
                run_count = paragraph.add_run(str(status_info.get("count", 0)))
                run_count.font.size = Pt(12)
                run_count.font.name = 'Times New Roman'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Получаем ключ для списка сотрудников
                name_field = status_info.get("name", "")
                if "(" in name_field:
                    # Например, "в командировке (business trip)" -> вычленяем "в командировке"
                    status_name_russian = name_field.split("(")[0].strip()
                    employee_key = f"сотрудники со статусом {status_name_russian}"
                else:
                    # Если нет скобок, берём всё как есть
                    employee_key = f"сотрудники со статусом {name_field.strip()}"

                # Получаем список сотрудников
                employees = status_info.get(employee_key, [])

                # Добавляем детальные сведения о каждом сотруднике
                for employee in employees:
                    employee_text = (
                        f"\n{employee['ФИО']}"
                        f"\n{employee['Дата начало']}-{employee['Дата окончание']}"
                        f"\n{employee['Примечание']}"
                    )
                    run_employee = paragraph.add_run(employee_text)
                    run_employee.font.size = Pt(5)  # Мелкий шрифт для деталей
                    run_employee.font.name = 'Times New Roman'

            # Если это последняя строка таблицы, делаем её жирной
            if row_idx == total_rows - 1:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

            row_idx += 1

        # Сохраняем результат в буфер
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


state_service = StateService(State)
